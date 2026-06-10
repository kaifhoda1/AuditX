from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import os
from datetime import datetime

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), hex_color)
    shd.set(qn('w:val'), 'clear')
    tcPr.append(shd)

def set_cell_border(cell, color="D1D5DB"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border in ['top', 'left', 'bottom', 'right']:
        b = OxmlElement(f'w:{border}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)

def extract_section(text, section_name):
    pattern = rf'{section_name}:\s*\n(.*?)(?=\n[A-Z\s]+:|$)'
    match = re.search(pattern, text, re.DOTALL)
    if not match: return []
    block = match.group(1)
    items = re.findall(r'-\s(.+)', block)
    return [item.strip() for item in items if item.strip()]

def generate_word_report(scored_results: dict, overall: dict, company_name: str) -> str:
    os.makedirs("outputs/reports", exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = company_name.replace(" ", "_").lower()
    filename = f"outputs/reports/{safe_name}_{timestamp}.docx"

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── HEADER TABLE ─────────────────────────────────────
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.columns[0].width = Inches(4)
    header_table.columns[1].width = Inches(2.5)

    left = header_table.cell(0, 0)
    right = header_table.cell(0, 1)

    set_cell_bg(left, '0D1117')
    set_cell_bg(right, '0D1117')

    p = left.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run('BYTEFORTIX SECURITY')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x00, 0xD4, 0xAA)
    run.font.name = 'Courier New'

    p2 = left.add_paragraph()
    r2 = p2.add_run('DPDP COMPLIANCE GAP ANALYSIS REPORT')
    r2.font.size = Pt(8)
    r2.font.color.rgb = RGBColor(0x94, 0xA3, 0xB8)
    r2.font.name = 'Courier New'

    p3 = left.add_paragraph()
    r3 = p3.add_run('Powered by AuditX — AI-Assisted Compliance Auditing')
    r3.font.size = Pt(7)
    r3.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
    r3.font.name = 'Courier New'
    r3.font.italic = True

    rp = right.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for line in ['kaifhoda1@gmail.com', 'github.com/kaifhoda1', 'bytefortix.com']:
        rr = rp.add_run(line + '\n')
        rr.font.size = Pt(7)
        rr.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
        rr.font.name = 'Courier New'

    doc.add_paragraph()

    # ── CLIENT INFO TABLE ─────────────────────────────────
    now = datetime.now().strftime("%d %B %Y")
    info_table = doc.add_table(rows=3, cols=4)
    info_table.style = 'Table Grid'
    info_data = [
        ['Client Company', company_name, 'Audit Date', now],
        ['Overall Score', f"{overall['score']}/100", 'Risk Level', overall['label']],
        ['Frameworks', str(len(scored_results)), 'Status', 'DRAFT — Awaiting Auditor Review'],
    ]
    for i, row_data in enumerate(info_data):
        row = info_table.rows[i]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = val
            p = cell.paragraphs[0]
            run = p.runs[0]
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
            if j % 2 == 0:
                run.bold = True
                set_cell_bg(cell, 'F4F5F7')
            set_cell_border(cell)

    doc.add_paragraph()

    # ── COMPLIANCE SCORE SUMMARY ──────────────────────────
    h = doc.add_paragraph()
    hr = h.add_run('COMPLIANCE SCORE SUMMARY')
    hr.bold = True
    hr.font.size = Pt(11)
    hr.font.color.rgb = RGBColor(0x0D, 0x11, 0x17)
    hr.font.name = 'Courier New'

    score_table = doc.add_table(rows=1 + len(scored_results) + 1, cols=4)
    score_table.style = 'Table Grid'

    headers = ['Framework', 'Score', 'Risk Level', 'Result']
    header_row = score_table.rows[0]
    for j, h in enumerate(headers):
        cell = header_row.cells[j]
        cell.text = h
        p = cell.paragraphs[0]
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name = 'Calibri'
        set_cell_bg(cell, '0D1117')

    # Overall row
    overall_row = score_table.rows[1]
    overall_data = ['OVERALL', f"{overall['score']}/100", overall['label'], '---']
    for j, val in enumerate(overall_data):
        cell = overall_row.cells[j]
        cell.text = val
        p = cell.paragraphs[0]
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
        set_cell_bg(cell, 'F0FDF4')
        set_cell_border(cell)

    risk_colors = {
        'LOW RISK': '166534',
        'MEDIUM RISK': '92400E',
        'HIGH RISK': '991B1B',
        'CRITICAL RISK': '7F1D1D'
    }

    for i, (fw, result) in enumerate(scored_results.items(), 2):
        row = score_table.rows[i]
        result_text = 'PASS' if result.get('score', 0) >= 60 else 'FAIL'
        row_data = [
            result.get('framework_name', fw),
            f"{result.get('score', 0)}/100",
            result.get('label', ''),
            result_text
        ]
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = val
            p = cell.paragraphs[0]
            run = p.runs[0]
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
            if j == 2:
                color = risk_colors.get(result.get('label', ''), '000000')
                r, g, b = int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
                run.font.color.rgb = RGBColor(r, g, b)
                run.bold = True
            set_cell_border(cell)

    doc.add_paragraph()

    # ── PER FRAMEWORK ─────────────────────────────────────
    for fw, result in scored_results.items():
        fw_name = result.get('framework_name', fw)
        score = result.get('score', 0)
        label = result.get('label', '')

        h = doc.add_paragraph()
        hr = h.add_run(f'SECTION — {fw_name.upper()}')
        hr.bold = True
        hr.font.size = Pt(10)
        hr.font.color.rgb = RGBColor(0x0D, 0x11, 0x17)
        hr.font.name = 'Courier New'

        sub = doc.add_paragraph()
        sr = sub.add_run(f'Score: {score}/100  |  Risk: {label}')
        sr.font.size = Pt(9)
        sr.font.color.rgb = RGBColor(0x47, 0x55, 0x69)
        sr.font.name = 'Calibri'
        sr.italic = True

        if 'error' in result:
            p = doc.add_paragraph(f"Error: {result['error']}")
            continue

        analysis = result.get('analysis', '')
        passed = extract_section(analysis, "PASSED CHECKS")
        failed = extract_section(analysis, "FAILED CHECKS")
        gaps = extract_section(analysis, "GAP ANALYSIS")
        fixes = extract_section(analysis, "RECOMMENDED FIXES")

        # Section 1
        s1 = doc.add_paragraph()
        s1r = s1.add_run('SECTION 1 — WHAT YOU ARE DOING RIGHT')
        s1r.bold = True
        s1r.font.size = Pt(9)
        s1r.font.color.rgb = RGBColor(0x16, 0x65, 0x34)
        s1r.font.name = 'Calibri'
        for item in passed:
            p = doc.add_paragraph(style='List Bullet')
            r = p.add_run(item)
            r.font.size = Pt(9)
            r.font.name = 'Calibri'

        # Section 2 — Gaps table
        s2 = doc.add_paragraph()
        s2r = s2.add_run('SECTION 2 — COMPLIANCE GAPS FOUND')
        s2r.bold = True
        s2r.font.size = Pt(9)
        s2r.font.color.rgb = RGBColor(0x99, 0x1B, 0x1B)
        s2r.font.name = 'Calibri'

        if failed:
            gap_table = doc.add_table(rows=1 + len(failed), cols=3)
            gap_table.style = 'Table Grid'
            for j, h in enumerate(['Gap Identified', 'Risk', 'Recommendation']):
                cell = gap_table.rows[0].cells[j]
                cell.text = h
                run = cell.paragraphs[0].runs[0]
                run.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.name = 'Calibri'
                set_cell_bg(cell, '0D1117')

            for i, (fail_item, fix_item) in enumerate(zip(failed, fixes if fixes else ['Review and update policy'] * len(failed)), 1):
                row = gap_table.rows[i]
                row.cells[0].text = fail_item
                row.cells[1].text = 'HIGH'
                row.cells[2].text = fix_item
                for j in range(3):
                    run = row.cells[j].paragraphs[0].runs[0]
                    run.font.size = Pt(8)
                    run.font.name = 'Calibri'
                    set_cell_border(row.cells[j])

        # Section 3 — Action plan
        s3 = doc.add_paragraph()
        s3r = s3.add_run('SECTION 3 — PRIORITY ACTION PLAN')
        s3r.bold = True
        s3r.font.size = Pt(9)
        s3r.font.color.rgb = RGBColor(0x0D, 0x11, 0x17)
        s3r.font.name = 'Calibri'

        if fixes:
            action_table = doc.add_table(rows=1 + len(fixes), cols=4)
            action_table.style = 'Table Grid'
            for j, h in enumerate(['#', 'Action Required', 'Timeline', 'Status']):
                cell = action_table.rows[0].cells[j]
                cell.text = h
                run = cell.paragraphs[0].runs[0]
                run.bold = True
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.name = 'Calibri'
                set_cell_bg(cell, '0D1117')

            for i, fix in enumerate(fixes, 1):
                row = action_table.rows[i]
                row.cells[0].text = str(i)
                row.cells[1].text = fix
                row.cells[2].text = 'Within 2 weeks'
                row.cells[3].text = 'Pending'
                for j in range(4):
                    run = row.cells[j].paragraphs[0].runs[0]
                    run.font.size = Pt(8)
                    run.font.name = 'Calibri'
                    set_cell_border(row.cells[j])

        doc.add_paragraph()

    # ── DISCLAIMER ───────────────────────────────────────
    disc_table = doc.add_table(rows=1, cols=1)
    disc_table.style = 'Table Grid'
    cell = disc_table.cell(0, 0)
    set_cell_bg(cell, 'FFFBEB')
    p = cell.paragraphs[0]
    r = p.add_run('DISCLAIMER: This report is DRAFT — Awaiting Auditor Review. '
                  'AuditX outputs are AI-assisted analysis only. Not legal advice. '
                  'Always verify with a qualified compliance professional. '
                  'ByteFortix Security accepts no liability for decisions made based on this report.')
    r.font.size = Pt(8)
    r.font.name = 'Calibri'
    r.font.color.rgb = RGBColor(0x78, 0x35, 0x0F)
    set_cell_border(cell, 'FBB024')

    doc.add_paragraph()

    # ── SIGNATURE ─────────────────────────────────────────
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.style = 'Table Grid'
    left = sig_table.cell(0, 0)
    right = sig_table.cell(0, 1)
    set_cell_border(left)
    set_cell_border(right)

    lp = left.paragraphs[0]
    lr = lp.add_run('Prepared by: Mohammad Kaif\nByteFortix Security\nkaifhoda1@gmail.com')
    lr.font.size = Pt(8)
    lr.font.name = 'Calibri'

    rp = right.paragraphs[0]
    rr = rp.add_run(f'Received by: _______________\n{company_name}\nDate: _______________')
    rr.font.size = Pt(8)
    rr.font.name = 'Calibri'

    doc.save(filename)
    return filename

if __name__ == "__main__":
    dummy = {
        "dpdp": {
            "framework_name": "DPDP Act 2023",
            "score": 70, "label": "MEDIUM RISK",
            "analysis": """COMPLIANCE SCORE: 70

PASSED CHECKS:
- Company collects data with stated purpose
- Users can request deletion

FAILED CHECKS:
- No consent mechanism mentioned
- No Data Protection Officer named

GAP ANALYSIS:
- Missing explicit consent collection at signup
- No DPO contact information provided

RECOMMENDED FIXES:
- Add consent checkbox at registration
- Appoint a Data Protection Officer"""
        }
    }
    overall = {"score": 70, "label": "MEDIUM RISK"}
    path = generate_word_report(dummy, overall, "Test Hospital Pvt Ltd")
    print(f"Word report saved to: {path}")
